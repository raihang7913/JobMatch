"""
CV Generator Module
Generates ATS-friendly CV files (PDF) from parsed CV data and optimizations
"""

from fpdf import FPDF
from datetime import datetime
from typing import Dict, List
import os
import re
import copy


class ATSFriendlyCV(FPDF):
    """ATS-friendly CV generator with clean formatting"""
    
    @staticmethod
    def _sanitize(text: str) -> str:
        """Sanitize text for Latin-1 encoding (Arial font compatibility)"""
        if not text:
            return text
        replacements = {
            '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'",
            '\u201c': '"', '\u201d': '"', '\u2022': '*', '\u2026': '...',
            '\u00a0': ' ', '\r': '', '\x00': '',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        text = text.encode('latin-1', errors='ignore').decode('latin-1')
        return text
    
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
    
    def cell(self, w, h=0, text='', *args, **kwargs):
        return super().cell(w, h, self._sanitize(text), *args, **kwargs)
    
    def multi_cell(self, w, h=0, text='', *args, **kwargs):
        return super().multi_cell(w, h, self._sanitize(text), *args, **kwargs)
    
    def header(self):
        """Page header - keep minimal for ATS"""
        pass
    
    def footer(self):
        """Page footer - keep minimal for ATS"""
        pass
    
    def add_section_title(self, title: str):
        """Add section title"""
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title.upper(), 0, 1, 'L')
        self.ln(2)
    
    def add_text(self, text: str, bold: bool = False):
        """Add regular text"""
        style = 'B' if bold else ''
        self.set_font('Arial', style, 10)
        self.multi_cell(0, 5, text)
        self.ln(1)
    
    def add_bullet_point(self, text: str):
        """Add bullet point"""
        self.set_font('Arial', '', 10)
        self.cell(10, 5, chr(149), 0, 0)  # Bullet character
        self.multi_cell(0, 5, text)


class CVGenerator:
    """Generate optimized, ATS-friendly CV"""
    
    @staticmethod
    def _sanitize_text(text: str) -> str:
        """Sanitize text for Latin-1 encoding (Arial font compatibility)"""
        if not text:
            return text
        
        # Replace common Unicode characters with ASCII equivalents
        replacements = {
            '\u2013': '-',  # En dash
            '\u2014': '-',  # Em dash
            '\u2018': "'",  # Left single quote
            '\u2019': "'",  # Right single quote
            '\u201c': '"',  # Left double quote
            '\u201d': '"',  # Right double quote
            '\u2022': '*',  # Bullet point
            '\u2026': '...', # Ellipsis
        }
        
        for unicode_char, ascii_char in replacements.items():
            text = text.replace(unicode_char, ascii_char)
        
        # Remove any remaining non-Latin-1 characters
        text = text.encode('latin-1', errors='ignore').decode('latin-1')
        
        return text
    
    def __init__(self, cv_data: Dict, optimization: Dict = None):
        """
        Initialize generator with CV data and optional optimization
        
        Args:
            cv_data: Parsed CV data from cv_parser
            optimization: Optimization suggestions from cv_optimizer (optional)
        """
        self.cv_data = cv_data
        self.optimization = optimization
        
    def generate_pdf(self, output_path: str, job_title: str = None) -> str:
        """
        Generate ATS-friendly PDF CV
        
        Args:
            output_path: Path to save PDF
            job_title: Optional job title for optimization context
            
        Returns:
            Path to generated PDF file
        """
        pdf = ATSFriendlyCV()
        pdf.add_page()
        
        # Use optimized content if available, otherwise use original
        if self.optimization and 'optimizations' in self.optimization:
            opts = self.optimization['optimizations']
            summary = opts.get('summary', self.cv_data.get('summary'))
            skills = self._get_optimized_skills()
        else:
            summary = self.cv_data.get('summary')
            skills = self.cv_data.get('skills', [])
        
        # Contact Information (top of page)
        self._add_contact_section(pdf)
        
        # Professional Summary
        if summary:
            self._add_summary_section(pdf, summary, job_title)
        
        # Skills
        if skills:
            self._add_skills_section(pdf, skills)
        
        # Experience
        if self.cv_data.get('experience'):
            self._add_experience_section(pdf)
        
        # Education
        if self.cv_data.get('education'):
            self._add_education_section(pdf)
        
        # Save PDF
        pdf.output(output_path)
        return output_path
    
    def _add_contact_section(self, pdf: ATSFriendlyCV):
        """Add contact information at top"""
        name = self.cv_data.get('name', 'YOUR NAME')
        email = self.cv_data.get('email', '')
        phone = self.cv_data.get('phone', '')
        location = self.cv_data.get('location', '')
        
        # Name - Large and bold
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, name, 0, 1, 'C')
        pdf.ln(2)
        
        # Contact details - Centered, smaller
        pdf.set_font('Arial', '', 10)
        contact_info = []
        if email:
            contact_info.append(email)
        if phone:
            contact_info.append(phone)
        if location:
            contact_info.append(location)
        
        contact_line = ' | '.join(contact_info)
        contact_line = self._sanitize_text(contact_line)
        pdf.cell(0, 5, contact_line, 0, 1, 'C')
        pdf.ln(5)
    
    def _add_summary_section(self, pdf: ATSFriendlyCV, summary: str, job_title: str = None):
        """Add professional summary"""
        title = f"PROFESSIONAL SUMMARY"
        if job_title:
            title = f"OBJECTIVE - {job_title.upper()}"
        
        pdf.add_section_title(title)
        pdf.add_text(summary)
        pdf.ln(3)
    
    def _get_optimized_skills(self) -> List[str]:
        """Get skills list, prioritized if optimization available"""
        original_skills = self.cv_data.get('skills', [])
        
        if not self.optimization:
            return original_skills
        
        # Get matched skills from analysis (these should be highlighted)
        analysis = self.optimization.get('analysis', {})
        matched = analysis.get('matched_skills', [])
        
        # Reorder: matched skills first, then others
        skills_lower = {s.lower(): s for s in original_skills}
        matched_lower = [s.lower() for s in matched]
        
        prioritized = []
        # Add matched skills first
        for skill in matched:
            skill_lower = skill.lower()
            if skill_lower in skills_lower:
                prioritized.append(skills_lower[skill_lower])
        
        # Add remaining skills
        for skill in original_skills:
            if skill not in prioritized:
                prioritized.append(skill)
        
        return prioritized
    
    def _add_skills_section(self, pdf: ATSFriendlyCV, skills: List[str]):
        """Add skills section - ATS friendly format"""
        pdf.add_section_title("TECHNICAL SKILLS")
        
        # Group skills by category (simple heuristic)
        languages = []
        frameworks = []
        tools = []
        other = []
        
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower in ['python', 'java', 'javascript', 'typescript', 'go', 'c++', 'c#', 'php', 'ruby', 'rust']:
                languages.append(skill)
            elif skill_lower in ['react', 'vue', 'angular', 'django', 'flask', 'spring', 'express']:
                frameworks.append(skill)
            elif skill_lower in ['git', 'github', 'docker', 'kubernetes', 'jenkins', 'aws', 'azure', 'gcp']:
                tools.append(skill)
            else:
                other.append(skill)
        
        # Print categorized skills
        if languages:
            pdf.add_text("Languages: " + ", ".join(languages), bold=False)
        if frameworks:
            pdf.add_text("Frameworks: " + ", ".join(frameworks), bold=False)
        if tools:
            pdf.add_text("Tools & Platforms: " + ", ".join(tools), bold=False)
        if other:
            pdf.add_text("Other: " + ", ".join(other), bold=False)
        
        pdf.ln(3)
    
    def _add_experience_section(self, pdf: ATSFriendlyCV):
        """Add work experience section"""
        pdf.add_section_title("PROFESSIONAL EXPERIENCE")
        
        experiences = self.cv_data.get('experience', [])
        
        for exp in experiences:
            title = exp.get('title', 'Position')
            dates = exp.get('dates', 'N/A')
            description = exp.get('description', '')
            
            # Position title and dates
            pdf.set_font('Arial', 'B', 10)
            pdf.cell(0, 5, title, 0, 1, 'L')
            
            pdf.set_font('Arial', 'I', 9)
            pdf.cell(0, 5, dates, 0, 1, 'L')
            pdf.ln(1)
            
            # Description - split into bullet points if possible
            if description:
                desc_lines = description.split('\n')
                for line in desc_lines:
                    line = line.strip()
                    if line:
                        if line.startswith(('-', '*', '•')):
                            line = line[1:].strip()
                        pdf.add_bullet_point(line)
            
            pdf.ln(2)
    

    def generate_docx(self, output_path: str, job_title: str = None) -> str:
        """Generate an editable ATS-friendly DOCX CV."""
        from docx import Document

        doc = Document()
        name = self.cv_data.get('name') or self.cv_data.get('personal_info', {}).get('name') or 'YOUR NAME'
        doc.add_heading(name, 0)

        contact = [self.cv_data.get('email'), self.cv_data.get('phone'), self.cv_data.get('location')]
        contact = [str(x) for x in contact if x]
        if contact:
            doc.add_paragraph(' | '.join(contact))

        if self.optimization and 'optimizations' in self.optimization:
            summary = self.optimization['optimizations'].get('summary') or self.cv_data.get('summary')
            skills = self._get_optimized_skills()
        else:
            summary = self.cv_data.get('summary')
            skills = self.cv_data.get('skills', [])

        if summary:
            doc.add_heading(f"Objective - {job_title}" if job_title else "Professional Summary", level=1)
            doc.add_paragraph(summary)

        if skills:
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph(', '.join(skills))

        experience = self.cv_data.get('experience', [])
        if experience:
            doc.add_heading("Professional Experience", level=1)
            for exp in experience:
                title = exp.get('title', 'Position')
                company = exp.get('company', '')
                dates = exp.get('dates', '')
                heading = f"{title} - {company}" if company else title
                if dates:
                    heading += f" ({dates})"
                doc.add_paragraph(heading)
                for line in str(exp.get('description', '')).split('\n'):
                    line = line.strip().lstrip('-*• ')
                    if line:
                        doc.add_paragraph(line, style='List Bullet')

        education = self.cv_data.get('education', [])
        if education:
            doc.add_heading("Education", level=1)
            for edu in education:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                year = edu.get('year') or edu.get('dates', '')
                doc.add_paragraph(' - '.join(str(x) for x in [degree, institution, year] if x))

        doc.save(output_path)
        return output_path

    def _add_education_section(self, pdf: ATSFriendlyCV):
        """Add education section"""
        pdf.add_section_title("EDUCATION")
        
        education = self.cv_data.get('education', [])
        
        for edu in education:
            degree = edu.get('degree', 'Degree')
            year = edu.get('year', 'N/A')
            
            pdf.set_font('Arial', 'B', 10)
            pdf.cell(0, 5, degree, 0, 1, 'L')
            
            pdf.set_font('Arial', '', 9)
            pdf.cell(0, 5, f"Graduated: {year}", 0, 1, 'L')
            pdf.ln(2)


def generate_optimized_cv(cv_data: Dict, job_title: str, output_dir: str = None) -> Dict:
    """
    Convenience function to generate optimized CV
    
    Args:
        cv_data: Parsed CV data
        job_title: Target job title
        output_dir: Directory to save CV (default: same as CV location)
        
    Returns:
        Dict with file path and status
    """
    from cv_optimizer import optimize_cv_for_job
    
    # Create job dict for optimization
    job = {
        'title': job_title,
        'company': 'Target Company',
        'summary': job_title
    }
    
    # Get optimization suggestions
    optimization = optimize_cv_for_job(cv_data, job)
    
    # Generate filename
    name = cv_data.get('name', 'CV')
    name_clean = name.replace(' ', '_').replace('.', '')
    timestamp = datetime.now().strftime('%Y%m%d')
    filename = f"{name_clean}_Optimized_{timestamp}.pdf"
    
    # Determine output directory
    if output_dir is None:
        output_dir = os.path.dirname(cv_data.get('cv_path', '.'))
    
    output_path = os.path.join(output_dir, filename)
    
    # Generate CV
    generator = CVGenerator(cv_data, optimization)
    result_path = generator.generate_pdf(output_path, job_title)
    
    return {
        'status': 'success',
        'file_path': result_path,
        'filename': filename,
        'optimization_applied': True,
        'job_title': job_title
    }


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace wording in a paragraph while keeping paragraph style/template position."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ''


def generate_tailored_docx_from_template(template_path: str, cv_data: Dict, applied_suggestions: Dict, output_path: str) -> str:
    """Edit wording inside the uploaded DOCX template; do not rebuild layout."""
    from docx import Document

    doc = Document(template_path)
    replacements = []
    summary = applied_suggestions.get('summary')
    if summary and cv_data.get('summary'):
        replacements.append((cv_data['summary'], summary))

    experiences = cv_data.get('experience', [])
    for item in applied_suggestions.get('experience', []):
        desc = item.get('description')
        if not desc:
            continue
        original = item.get('original')
        if not original:
            try:
                original = experiences[int(item.get('index'))].get('description')
            except (TypeError, ValueError, IndexError):
                original = None
        if original:
            replacements.append((original, desc))

    changed = 0
    for old, new in replacements:
        old = str(old).strip()
        if not old:
            continue
        candidates = [old] + [line.strip() for line in old.split('\n') if len(line.strip()) > 20]
        found = False
        for candidate in candidates:
            for paragraph in doc.paragraphs:
                if candidate in paragraph.text:
                    _replace_paragraph_text(paragraph, paragraph.text.replace(candidate, new))
                    changed += 1
                    found = True
                    break
            if found:
                break

    # ponytail: if the source DOCX has no summary paragraph, append one instead of inventing layout matching.
    if summary and not cv_data.get('summary'):
        doc.add_paragraph(summary)
        changed += 1

    if changed == 0:
        raise ValueError('No matching text found in DOCX template to replace')

    doc.save(output_path)
    return output_path

def build_tailored_cv_data(cv_data: Dict, applied_suggestions: Dict) -> Dict:
    """Merge selected AI suggestions into parsed CV data without mutating original."""
    tailored = copy.deepcopy(cv_data)

    summary = applied_suggestions.get('summary')
    if summary:
        tailored['summary'] = summary

    skills = applied_suggestions.get('skills')
    if skills:
        seen = set()
        tailored['skills'] = [s for s in skills if s and not (s.lower() in seen or seen.add(s.lower()))]

    experiences = tailored.get('experience', [])
    for item in applied_suggestions.get('experience', []):
        try:
            idx = int(item.get('index'))
        except (TypeError, ValueError):
            continue
        if 0 <= idx < len(experiences) and item.get('description'):
            experiences[idx]['description'] = item['description']

    return tailored


def generate_tailored_cv(cv_data: Dict, applied_suggestions: Dict, output_dir: str, job_title: str = '', output_format: str = 'pdf', template_path: str = '') -> Dict:
    """Generate tailored CV. DOCX templates are edited in-place wording-only when possible."""
    os.makedirs(output_dir, exist_ok=True)
    fmt = output_format.lower()
    if fmt not in ('pdf', 'docx'):
        raise ValueError('format must be pdf or docx')

    tailored = build_tailored_cv_data(cv_data, applied_suggestions)
    name = tailored.get('name') or tailored.get('personal_info', {}).get('name') or 'CV'
    safe_name = re.sub(r'[^A-Za-z0-9_-]+', '_', name).strip('_') or 'CV'
    safe_job = re.sub(r'[^A-Za-z0-9_-]+', '_', job_title or 'tailored').strip('_') or 'tailored'
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{safe_name}_{safe_job}_tailored_{timestamp}.{fmt}"
    output_path = os.path.join(output_dir, filename)

    optimization = {'optimizations': {'summary': tailored.get('summary')}}
    generator = CVGenerator(tailored, optimization)
    if fmt == 'docx' and template_path and str(template_path).lower().endswith('.docx'):
        generate_tailored_docx_from_template(template_path, cv_data, applied_suggestions, output_path)
    elif fmt == 'docx':
        generator.generate_docx(output_path, job_title)
    else:
        generator.generate_pdf(output_path, job_title)

    return {'status': 'success', 'file_path': output_path, 'filename': filename, 'format': fmt}
